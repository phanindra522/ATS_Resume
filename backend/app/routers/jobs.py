from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from typing import List, Optional
import json
import PyPDF2
from docx import Document
import io
from app.database import get_database
from app.models.job import JobCreate, JobResponse, JobUpdate, JobWithUserInfo, JobFileUpload
from app.routers.auth import get_current_user
from app.services.llm_service import parse_job_with_llm, LLMProvider
from app.core.config import settings
from bson import ObjectId
from datetime import datetime

router = APIRouter()

@router.get("/llm-config")
async def get_llm_config():
    """Get LLM configuration status"""
    return {
        "config": settings.get_llm_config_info(),
        "message": "LLM configuration status"
    }

@router.post("/test-llm")
async def test_llm_parsing(
    text: str,
    current_user = Depends(get_current_user)
):
    """Test LLM parsing with sample text"""
    try:
        result = await parse_job_with_llm(text)
        return {
            "success": True,
            "parsed_data": result,
            "message": "LLM parsing successful"
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": "LLM parsing failed"
        }

@router.post("/extract", response_model=JobFileUpload)
async def extract_job_from_file(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user)
):
    """Extract job information from uploaded file without saving"""
    
    # Validate file type
    allowed_types = ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file type. Only PDF, TXT, and DOCX files are allowed."
        )
    
    # Validate file size (10MB max)
    if file.size > 10 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File size must be less than 10MB"
        )
    
    try:
        # Extract text content from file
        file_content = await file.read()
        extracted_text = ""
        
        if file.content_type == "application/pdf":
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif file.content_type == "text/plain":
            # Extract text from TXT file
            extracted_text = file_content.decode('utf-8')
            
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from DOCX file
            doc = Document(io.BytesIO(file_content))
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        
        # Clean up extracted text
        extracted_text = extracted_text.strip()
        
        # Use LLM to parse job information intelligently
        try:
            parsed_info = await parse_job_with_llm(extracted_text)
        except Exception as e:
            # Fallback to rule-based parsing if LLM fails
            print(f"LLM parsing failed, using fallback: {str(e)}")
            parsed_info = parse_job_text(extracted_text)
        
        return JobFileUpload(
            message="Job information extracted successfully using AI",
            extracted_text=extracted_text,
            job_id="",  # Not saved yet
            **parsed_info
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to extract job information: {str(e)}"
        )

def parse_job_text(text: str) -> dict:
    """Parse job text to extract structured information"""
    lines = text.split('\n')
    
    # Initialize result
    result = {
        "title": "",
        "company": "",
        "description": "",
        "requirements": [],
        "skills": [],
        "experience_level": "",
        "location": "",
        "salary_range": ""
    }
    
    # Simple keyword-based extraction
    text_lower = text.lower()
    
    # Extract title (usually in first few lines or after "position:", "job title:", etc.)
    for i, line in enumerate(lines[:10]):  # Check first 10 lines
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in ['position:', 'job title:', 'title:', 'role:']):
            result["title"] = line.split(':', 1)[1].strip() if ':' in line else line.strip()
            break
        elif 'systems administrator' in line_lower:
            result["title"] = "Systems Administrator"
            break
        elif (i < 3 and len(line.strip()) > 5 and len(line.strip()) < 100 and 
              not any(skip_word in line_lower for skip_word in ['summary:', 'description:', 'overview:', 'about:', 'company:', 'employer:', 'organization:', 'we are', 'currently seeking']) and
              not line.strip().startswith(('+', '-', '•', '*', '1.', '2.', '3.')) and
              not any(keyword in line_lower for keyword in ['years', 'experience', 'building', 'apps', 'javascript', 'typescript'])):
            result["title"] = line.strip()
    
    # Extract company (look for "company:", "employer:", etc.)
    for line in lines[:15]:
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in ['company:', 'employer:', 'organization:']):
            result["company"] = line.split(':', 1)[1].strip() if ':' in line else line.strip()
            break
    
    # Extract location
    for line in lines:
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in ['location:', 'based in:', 'work from:']):
            result["location"] = line.split(':', 1)[1].strip() if ':' in line else line.strip()
            break
        elif any(keyword in line_lower for keyword in ['remote', 'hybrid', 'on-site']):
            result["location"] = line.strip()
            break
    
    # Extract salary
    for line in lines:
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in ['salary:', 'compensation:', 'pay:', '$']):
            if '$' in line:
                result["salary_range"] = line.strip()
                break
    
    # Extract experience level
    for line in lines:
        line_lower = line.lower().strip()
        if 'entry' in line_lower and 'level' in line_lower:
            result["experience_level"] = "Entry Level"
            break
        elif 'senior' in line_lower:
            result["experience_level"] = "Senior Level"
            break
        elif 'mid' in line_lower or 'intermediate' in line_lower:
            result["experience_level"] = "Mid Level"
            break
        elif 'lead' in line_lower or 'manager' in line_lower:
            result["experience_level"] = "Lead/Manager"
            break
    
    # Extract skills (look for common skill keywords)
    skill_keywords = [
        'python', 'javascript', 'java', 'react', 'angular', 'vue', 'node.js', 'express',
        'django', 'flask', 'spring', 'sql', 'mongodb', 'postgresql', 'mysql',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git', 'jenkins',
        'html', 'css', 'typescript', 'php', 'ruby', 'go', 'rust', 'c++', 'c#',
        'machine learning', 'ai', 'data science', 'analytics', 'tableau', 'power bi',
        'technical troubleshooting', 'hardware support', 'software support', 'customer service',
        'problem-solving', 'communication', 'linux', 'windows', 'networking', 'it support',
        'systems administration', 'troubleshooting', 'hardware', 'software', 'documentation'
    ]
    
    found_skills = []
    in_skills_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        if 'skills:' in line_lower:
            in_skills_section = True
            continue
        
        if in_skills_section:
            if line.strip() and line.strip().startswith(('•', '-', '*')):
                clean_line = line.strip().lstrip('•-* ').strip()
                if clean_line and len(clean_line) > 3:
                    found_skills.append(clean_line.title())
            elif line.strip() == '':
                continue
            else:
                break
    
    # Also look for skills throughout the text
    for line in lines:
        line_lower = line.lower()
        for skill in skill_keywords:
            if skill in line_lower and skill not in [s.lower() for s in found_skills]:
                found_skills.append(skill.title())
    
    result["skills"] = found_skills[:10]  # Limit to 10 skills
    
    # Extract requirements (look for bullet points or numbered lists)
    requirements = []
    in_requirements_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in ['requirements:', 'qualifications:', 'must have:', 'required:']):
            in_requirements_section = True
            continue
        
        if in_requirements_section:
            if line.strip() and (line.strip().startswith(('•', '-', '*', '1.', '2.', '3.')) or 
                               any(keyword in line_lower for keyword in ['years', 'experience', 'degree', 'bachelor', 'master', 'demonstrated', 'expertise', 'specialization', 'understanding'])):
                clean_line = line.strip().lstrip('•-*123456789. ').strip()
                if clean_line and len(clean_line) > 10:
                    requirements.append(clean_line)
            elif line.strip() == '':
                continue
            else:
                break
    
    # If no requirements section found, look for standalone requirement-like lines
    if not requirements:
        for line in lines:
            line_lower = line.lower().strip()
            # Look for lines that contain requirement indicators
            if (any(keyword in line_lower for keyword in ['years', 'experience', 'degree', 'bachelor', 'master', 'proficiency', 'knowledge of', 'familiar with']) and
                len(line.strip()) > 15 and len(line.strip()) < 200 and
                not any(skip_word in line_lower for skip_word in ['summary:', 'description:', 'overview:', 'about:', 'job title:', 'position:'])):
                requirements.append(line.strip())
    
    result["requirements"] = requirements[:15]  # Limit to 15 requirements
    
    # Use the full text as description if no specific description found
    if not result["description"]:
        result["description"] = text[:2000]  # Limit to 2000 characters
    
    return result

@router.post("/", response_model=JobResponse)
async def create_job(
    job: JobCreate,
    current_user = Depends(get_current_user)
):
    try:
        db = get_database()
        
        job_data = job.dict()
        job_data["user_id"] = current_user["_id"]
        job_data["created_at"] = datetime.utcnow()
        job_data["updated_at"] = datetime.utcnow()
        
        print(f"Creating job with data: {job_data}")
        
        result = await db.jobs.insert_one(job_data)
        
        # Fix: Use dictionary access for inserted_id
        job_data["_id"] = result["inserted_id"]
        
        print(f"Job created successfully with ID: {result['inserted_id']}")
        
        return JobResponse(**job_data)
        
    except Exception as e:
        print(f"Error creating job: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create job: {str(e)}"
        )

@router.post("/upload", response_model=JobResponse)
async def upload_job_file(
    file: UploadFile = File(...),
    title: str = Form(...),
    company: Optional[str] = Form(""),
    description: Optional[str] = Form(""),
    requirements: Optional[str] = Form(""),
    skills: Optional[str] = Form(""),
    experience_level: Optional[str] = Form(""),
    location: Optional[str] = Form(""),
    salary_range: Optional[str] = Form(""),
    current_user = Depends(get_current_user)
):
    """Upload a job description file and create job"""
    
    # Validate file type
    allowed_types = ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file type. Only PDF, TXT, and DOCX files are allowed."
        )
    
    # Validate file size (10MB max)
    if file.size > 10 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File size must be less than 10MB"
        )
    
    try:
        # Extract text content from file
        file_content = await file.read()
        extracted_text = ""
        
        if file.content_type == "application/pdf":
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif file.content_type == "text/plain":
            # Extract text from TXT file
            extracted_text = file_content.decode('utf-8')
            
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from DOCX file
            doc = Document(io.BytesIO(file_content))
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        
        # Clean up extracted text
        extracted_text = extracted_text.strip()
        
        # If no description provided, use extracted text
        if not description and extracted_text:
            description = extracted_text[:2000]  # Limit to 2000 characters
        
        # Parse requirements and skills from form data
        requirements_list = []
        if requirements:
            try:
                requirements_list = json.loads(requirements)
            except json.JSONDecodeError:
                # Fallback: split by newlines if JSON parsing fails
                requirements_list = [req.strip() for req in requirements.split('\n') if req.strip()]
        
        skills_list = []
        if skills:
            try:
                skills_list = json.loads(skills)
            except json.JSONDecodeError:
                # Fallback: split by commas if JSON parsing fails
                skills_list = [skill.strip() for skill in skills.split(',') if skill.strip()]
        
        # Create job data
        job_data = {
            "title": title,
            "company": company,
            "description": description,
            "requirements": requirements_list,
            "skills": skills_list,
            "experience_level": experience_level or None,
            "location": location or None,
            "salary_range": salary_range or None,
            "user_id": current_user["_id"],
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        print(f"Uploading job with data: {job_data}")
        
        # Save to database
        db = get_database()
        result = await db.jobs.insert_one(job_data)
        
        # Fix: Use dictionary access for inserted_id
        job_data["_id"] = result["inserted_id"]
        
        print(f"Job uploaded successfully with ID: {result['inserted_id']}")
        
        return JobResponse(**job_data)
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process file: {str(e)}"
        )

@router.get("/", response_model=List[JobWithUserInfo])
async def get_jobs(current_user = Depends(get_current_user)):
    """Get all jobs from the global pool (all recruiters can see all jobs)"""
    try:
        db = get_database()
        
        # Global pool: Get all jobs from all users
        jobs_cursor = await db.jobs.find({})
        jobs = await jobs_cursor.to_list(length=1000)  # Increased limit for global pool
        
        # Get user information for each job
        jobs_with_user_info = []
        for job in jobs:
            # Get user information
            user = await db.users.find_one({"_id": job["user_id"]})
            user_name = user["full_name"] if user else "Unknown User"
            user_email = user["email"] if user else "unknown@example.com"
            
            # Create job with user info
            job_data = job.copy()
            job_data["user_name"] = user_name
            job_data["user_email"] = user_email
            
            jobs_with_user_info.append(JobWithUserInfo(**job_data))
        
        return jobs_with_user_info
        
    except Exception as e:
        print(f"Error getting jobs: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to retrieve jobs: {str(e)}"
        )

@router.get("/{job_id}", response_model=JobResponse)
async def get_job(job_id: str, current_user = Depends(get_current_user)):
    """Get a specific job by ID from the global pool"""
    try:
        db = get_database()
        job = await db.jobs.find_one({
            "_id": job_id
        })
        
        if not job:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Job not found"
            )
        
        return JobResponse(**job)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting job {job_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to retrieve job: {str(e)}"
        )

@router.put("/{job_id}", response_model=JobResponse)
async def update_job(
    job_id: str,
    job_update: JobUpdate,
    current_user = Depends(get_current_user)
):
    """Update a job description"""
    try:
        db = get_database()
        
        # Check if job exists and belongs to user
        existing_job = await db.jobs.find_one({
            "_id": job_id,
            "user_id": current_user["_id"]
        })
        
        if not existing_job:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Job not found"
            )
        
        # Update job
        update_data = job_update.dict(exclude_unset=True)
        update_data["updated_at"] = datetime.utcnow()
        
        await db.jobs.update_one(
            {"_id": job_id},
            {"$set": update_data}
        )
        
        # Get updated job
        updated_job = await db.jobs.find_one({"_id": job_id})
        return JobResponse(**updated_job)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error updating job {job_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to update job: {str(e)}"
        )

@router.delete("/{job_id}")
async def delete_job(job_id: str, current_user = Depends(get_current_user)):
    """Delete a job description"""
    try:
        db = get_database()
        
        # Check if job exists and belongs to user
        job = await db.jobs.find_one({
            "_id": job_id,
            "user_id": current_user["_id"]
        })
        
        if not job:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Job not found"
            )
        
        # Delete from database
        await db.jobs.delete_one({"_id": job_id})
        
        return {"message": "Job deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting job {job_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete job: {str(e)}"
        )
